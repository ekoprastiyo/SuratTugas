
import streamlit as st

st.set_page_config(layout='wide') # layout='wide', Set the page layout to wide

import sys
sys.path.append('/content/drive/My Drive/Python/Packages/')
import pandas as pd
import numpy as np
from io import BytesIO
import requests
from datetime import datetime, timedelta
from docxtpl import DocxTemplate
from docx2pdf import convert
import jinja2
from docx import Document
import os
from docxcompose.composer import Composer
from pypdf import PdfWriter
from num2words import num2words # Add num2words import

# --- Configuration --- #
template_file = 'ST26-template.docx'
# template form absensi
temp_file_konfirmasi_absen = 'Konfirmasi_absen.docx'
temp_file_perintah_lembur = 'Surat Perintah Lembur.docx'
temp_file_daftar_lembur = 'Daftar Perintah Lembur.docx'
output_file = 'combined_output_2026.docx'
url = "https://docs.google.com/spreadsheets/d/1SORCi_jXxEN-HSXWBOjX19FY8a6FzegPSAPbuh5k1sI/export?format=xlsx"

# --- Data Loading and Preprocessing --- #
@st.cache_data(ttl=3600) # Cache data for 1 hour by default
def load_data():
    kegiatan_df = pd.read_excel(url, sheet_name='Kegiatan', header=2)
    # Vendor EPS dibuatkan Lembar Absensi
    Karyawan = pd.read_excel(url, sheet_name='Karyawan')
    karyawan_df = pd.read_excel(url, sheet_name='Karyawan_ST')
    kegiatan_df['TanggalAwal'] = pd.to_datetime(kegiatan_df['TanggalAwal'], format='%d/%m/%Y', errors='coerce')
    kegiatan_df['TanggalAkhir'] = pd.to_datetime(kegiatan_df['TanggalAkhir'], format='%d/%m/%Y', errors='coerce') # Corrected typo here
    # Filter out rows where 'Karyawan' is '-'
    kegiatan_df = kegiatan_df[kegiatan_df['Karyawan'] != '-']
    return kegiatan_df, Karyawan, karyawan_df

# Call load_data once. It will use cache if available, or fetch fresh data if cache cleared by the button.
Kegiatan, Karyawan, karyawan_df = load_data()

# --- CSS ---#
st.markdown(
    """
    <style>
    /* Remove padding from the main content container */
    .block-container {
        padding-top: 3rem;
        padding-bottom: 0rem;
        margin-top: -2rem;
    }
    /* Optional: Remove padding from the header area */

    # Reduce padding at the top of the sidebar
    [data-testid="stSidebarUserContent"] {
        padding-top: 1rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- Streamlit UI Components --- #
# --- Sidebar Content --- #
st.sidebar.header('Filter Options')

options = ["All", "Selesai", "Belum Selesai"]
Status_surat = st.sidebar.segmented_control(
    "Status Surat Tugas", options, selection_mode="single", default="All"
)

months = {
    'All': 0,
    'January': 1, 'February': 2, 'March': 3, 'April': 4, 'May': 5, 'June': 6,
    'July': 7, 'August': 8, 'September': 9, 'October': 10, 'November': 11, 'December': 12
}

current_month_name = datetime.now().strftime('%B')
try:
    default_month_index = list(months.keys()).index(current_month_name)
except ValueError:
    default_month_index = months[current_month_name]

selected_month_name = st.sidebar.selectbox('Select Month', list(months.keys()), index=default_month_index)
selected_month_num = months[selected_month_name]

list_name = Karyawan['Nama'].to_list()
list_name.insert(0, "All")
search_name = st.sidebar.selectbox('Search by Employee Name', list_name)

filtered_Kegiatan = Kegiatan.copy()

if Status_surat != "All":
    filtered_Kegiatan = filtered_Kegiatan[filtered_Kegiatan['Keterangan'] == Status_surat]

if selected_month_num != 0:
    filtered_Kegiatan = filtered_Kegiatan[filtered_Kegiatan['TanggalAwal'].dt.month == selected_month_num]

if search_name != "All":
    filtered_Kegiatan = filtered_Kegiatan[filtered_Kegiatan['Karyawan'].str.contains(search_name, case=False, na=False)]

# 1. Create a button in the app
if st.sidebar.button("Clear App Cache"):
    # 2. Clear both data and function caches
    st.cache_data.clear()
    st.cache_resource.clear()
    
    # 3. Show success message and rerun to refresh the page
    st.success("Cache cleared successfully!")
    st.rerun()

# --- Main Content --- #
st.title('Surat Tugas Kanwil X')
col1, col2 = st.columns([0.55, 0.45]) # Removed the third column
with col1:
  st.write('Displaying the DataFrame:')
with col2:
  st.markdown(f"<div style='text-align: right;'><a href='https://docs.google.com/spreadsheets/d/1SORCi_jXxEN-HSXWBOjX19FY8a6FzegPSAPbuh5k1sI/' target='_blank'>Go to Spreadsheet</a></div>", unsafe_allow_html=True)

filtered_Kegiatan_display = filtered_Kegiatan.drop(columns=['TanggalAwal', 'TanggalAkhir'], errors='ignore')
st.dataframe(
    filtered_Kegiatan_display, 
    width="stretch",
    column_config={
        "NoSurat": st.column_config.Column(width=None),
        "Keterangan": st.column_config.Column(width=None),
        "Karyawan": st.column_config.Column(width=None),
        "Lokasi": st.column_config.Column(width=None),
        "JangkaWaktu": st.column_config.Column(width=None),
        "Kegiatan": st.column_config.Column(width=None),
    }
    )

# --- Document Generation Logic --- #
def merge_docx_files(master_path, files_to_append, output_path):
    master = Document(master_path)
    master.add_page_break()
    composer = Composer(master)

    # for file_path in files_to_append:
    #     doc_to_append = Document(file_path)
    #     doc_to_append.add_page_break()
    #     composer.append(doc_to_append)
    for i, file_path in enumerate(files_to_append):
      # Load each additional document
      doc_to_append = Document(file_path)
      if i == 0:
        doc_to_append.add_page_break()
      # Append it to the master while preserving formatting
      composer.append(doc_to_append)
    composer.save(output_path)

def angka_ke_bulan(angka):
  bulan = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
    "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
  return bulan[angka]

def tanggal_eng_to_ina(date):
  hari = {"Monday":"Senin", "Tuesday":"Selasa", "Wednesday":"Rabu",
           "Thursday":"Kamis", "Friday":"Jumat", "Saturday":"Sabtu",
           "Sunday":"Minggu"
  }
  return hari.get(date, date)

def bulan_eng_to_ina(english_month_name):
  month_mapping = {
      "January": "Januari", "February": "Februari", "March": "Maret",
      "April": "April", "May": "Mei", "June": "Juni",
      "July": "Juli", "August": "Agustus", "September": "September",
      "October": "Oktober", "November": "November", "December": "Desember"
  }
  return month_mapping.get(english_month_name, english_month_name) # Return original if not found

# --- Download Surat Tugas --- #
def download_ST(df, Kegiatan_all, Karyawan_all):
    ListST = list(df["NoSurat"].unique())
    list_dicts = []

    for i in ListST:
        KegiatanPerST = Kegiatan_all[Kegiatan_all['NoSurat'] == int(i)]
        KaryawanPerST = Karyawan_all[Karyawan_all['NoSurat'] == int(i)]
        nama_karyawan = KaryawanPerST[["Nama", "NIK", "UnitKerja"]]
        kegiatan = KegiatanPerST[["JangkaWaktu", "Kegiatan", "Lokasi"]]

        nama_dict = nama_karyawan.to_dict(orient='records')
        kegiatan_dict = kegiatan.to_dict(orient='records')

        context = {
            "NoSurat": f"{i}",
            "table1": nama_dict,
            "table2": kegiatan_dict,
        }
        list_dicts.append(context)

    temp_files = []
    for i, entry in enumerate(list_dicts):
        tpl = DocxTemplate(template_file)
        tpl.render(entry)
        temp_name = f'temp_{i}.docx'
        tpl.save(temp_name)
        temp_files.append(temp_name)

    if list_dicts:
        merge_docx_files(temp_files[0], temp_files[1:], output_file)
        with open(output_file, "rb") as file:
            st.download_button(
                label="Download Surat Tugas",
                data=file,
                file_name=output_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("No documents generated for merging. Check your filters.")

    for f in temp_files:
        os.remove(f)


# --- Download Surat Lembur --- #

def Surat_Konfirmasi_absen(df, nama, NIK_Karyawan, bulan, temp_file_konfirmasi_absen):
  '''
  fungsi untuk menghapus data duplikat berdasarkan tanggal yang sama
  halaman konfirmasi absen perlu data tanggal lembur yang berbeda
  '''
  df_filtered = df.drop_duplicates(subset=["JangkaWaktu"])
  # Template DOCX
  # doc = DocxTemplate(template_file)
  TglBuat = datetime.now().strftime("%d %B %Y")
  TglBuat = TglBuat[:3] + bulan_eng_to_ina(TglBuat[3:-5]) + TglBuat[-5:]
  dict_lembur = {
      "JangkaWaktu" : "; ".join(list(df_filtered['JangkaWaktu'])),
      "TanggalBuat" : TglBuat,
      "nama" : nama,
      "NIK_Karyawan" : NIK_Karyawan,
  }

  temp_files = []
  output_file = 'konfirmasi_absen.docx'
  tpl = DocxTemplate(temp_file_konfirmasi_absen)
  tpl.render(dict_lembur)
  tpl.save(output_file)
  temp_files.append(output_file)
  return temp_files

def tgl_lembur_di_surat_perintah(list_tgl):
  if not list_tgl:
      return "", "", "", "" # Return empty strings if the list is empty

  list_tgl.sort()
  # tanggal awal surat tugas berhari-hari
  tgl_awal = datetime.strptime(list_tgl[0], "%d %b %Y")
  tgl_awal1 = tanggal_eng_to_ina(tgl_awal.strftime("%A"))
  tgl_awal2 = tgl_awal.strftime("%d %B %Y")
  # tanggal + nama bulan + tahun  dari variable tgl_awal2
  tgl_awal3 = tgl_awal2[:3] + bulan_eng_to_ina(tgl_awal2[3:-5]) + tgl_awal2[-5:]
  if len(list_tgl) > 1:
    # tanggal akhir surat tugas  berhari-hari
    tgl_akhir = datetime.strptime(list_tgl[-1], "%d %b %Y") # Changed to list[-1] to get the last date after sort
    tgl_akhir1 = tanggal_eng_to_ina(tgl_akhir.strftime("%A"))
    tgl_akhir2 = tgl_akhir.strftime("%d %B %Y")
    # tanggal + nama bulan + tahun  dari variable tgl_akhir2
    tgl_akhir3 = tgl_akhir2[:3] + bulan_eng_to_ina(tgl_akhir2[3:-5]) + tgl_akhir2[-5:]
  else:
    tgl_akhir1 = ""
    tgl_akhir3 = ""
  return tgl_awal1, tgl_awal3, tgl_akhir1, tgl_akhir3

def Surat_Perintah_Lembur(df, nama, bulan, temp_file_perintah_lembur):

  # Filter nomor surat
  ListST = list(df["NoSurat"].unique())

  # List dictionary
  list_dicts = []

  # def Surat_Perintah_Lembur(nama, NIK_Karyawan, bulan, output_file):
  for i in ListST:
    #filter kegiatan lembur per satu surat tugas
    KegiatanPerST = df[df['NoSurat'] == int(i)].reset_index(drop=True)

    #data kegiatan yang diinput dalam surat tugas
    # 1. Nama Kegiatan
    kegiatan = KegiatanPerST[["Kegiatan"]]

    # 2. No. Nama Kegiatan
    # tambah satu kolom nomor baru
    kegiatan = kegiatan.reset_index().rename(columns={'index': 'No'})
    kegiatan['No'] = kegiatan['No'] + 1

    # menggabungkan data 1 & 2
    # Dictionary lembur
    kegiatan_dict = kegiatan.to_dict(orient='records')
    # print(kegiatan_dict)

    # 3. Jumlah orang saat lembur
    # menghitung jumlah orang per surat tugas
    JumlahOrg = len(KegiatanPerST["Karyawan"].iloc[0].split(","))
    JumlahOrang = f'{JumlahOrg} ({num2words(JumlahOrg, lang='id')})'

    # 4. Jangka Waktu Lembur
    # list tanggal lembur
    # memastikan semua tanggal antara tanggal awal dan akhir terdata
    TanggalKegiatanPerST = []
    ST1hari = []
    STberharihari = []
    for x in range(len(KegiatanPerST)):
      # Use .iloc[x] for positional indexing
      # jika terdapat data di tanggal akhir maka lembur lebih dari satu hari
      # Changed check to pd.isna
      if pd.isna(KegiatanPerST['TanggalAkhir'].iloc[x]):
        ST1hari.append(KegiatanPerST['TanggalAwal'].iloc[x].strftime("%d %b %Y"))
      else:
        current_date = KegiatanPerST['TanggalAwal'].iloc[x]
        end_date = KegiatanPerST['TanggalAkhir'].iloc[x]
        delta = timedelta(days=1)

        # ketika tanggal awal belum sama dengan tanggal akhir
        # setiap tanggal ditambahkan ke dalam list
        # ditambah lagi satu hari hingga sama dengan end_date
        while current_date <= end_date:
            STberharihari.append(current_date.strftime("%d %b %Y"))
            current_date += delta

    # jika jumlah hari lembur lebih dari satu hari
    # maka dibuat teks tanggal lembur dari dan sampai dengan

    if len(list(set(STberharihari))) > 1:
      tgl_awal1, tgl_awal3, tgl_akhir1, tgl_akhir3 = tgl_lembur_di_surat_perintah(STberharihari)
      # teks jangka waktu surat tugas berhari - hari
      JangkaWaktu = f'{tgl_awal1} - {tgl_akhir1}, {tgl_awal3[:2]} - {tgl_akhir3}'
    else:
      tgl_awal1, tgl_awal3, tgl_akhir1, tgl_akhir3 = tgl_lembur_di_surat_perintah(ST1hari)

      if len(ST1hari) > 1:
        # teks jangka waktu surat tugas berhari - hari
        JangkaWaktu = f'{tgl_awal1} &amp; {tgl_akhir1}, {tgl_awal3[:2]} &amp; {tgl_akhir3}'
      else:
        JangkaWaktu = f'{tgl_awal1}, {tgl_awal3}'

    # 5. Hari Lembur *Weekend atau Weekdays
    # membuat teks weekend atau weekday sesuai hari lembur
    WaktuHari = []
    AlasanLembur = []
    # memastikan tidak ada tanggal yang sama
    TanggalKegiatanPerST = list(set(ST1hari + STberharihari))
    # looping setiap tanggal yang ada
    for y in TanggalKegiatanPerST:
      # ubah teks tanggal menjadi format tanggal
      current_date_obj = datetime.strptime(y, "%d %b %Y")
      # jika tanggal .weekday() kurang dari atau sama dengan lima artinya hari kerja (weekday)
      # jika tanggal .weekday() lebih dari 5 artinya akhir pekan (weekend)
      if current_date_obj.weekday() < 5:
        WaktuHari.append("Weekday")
        AlasanLembur.append("hingga malam hari")
      else:
        WaktuHari.append("Weekend")
        AlasanLembur.append("pada hari libur")

    # memastikan jika terdapat weekend dan weekday di satu surat tugas
    if "Weekend" in WaktuHari and "Weekday" in WaktuHari:
      WaktuHari = "(Weekend &amp; Weekday)"
      AlasanLembur = "pada hari libur &amp; hingga malam hari"
    elif "Weekend" in WaktuHari:
      WaktuHari = "(Weekend)"
      AlasanLembur = "pada hari libur"
    elif "Weekday" in WaktuHari:
      WaktuHari = "(Weekday)"
      AlasanLembur = "hingga malam hari"
    else:
      WaktuHari = ""
      AlasanLembur = ""

    # semua data dikumpul dalam satu dictionary
    # per dictionary diinput ke dalam list
    context = {
        "KegiatanLembur": kegiatan_dict,
        "JumlahOrang": JumlahOrang,
        "JangkaWaktu" : JangkaWaktu,
        "WaktuHari" : WaktuHari,
        "AlasanLembur" : AlasanLembur
            }
    list_dicts.append(context)

  temp_files = []
  # 2. Render each entry into a temporary file
  for i, entry in enumerate(list_dicts):
      tpl = DocxTemplate(temp_file_perintah_lembur)
      tpl.render(entry)
      temp_name = f'temp_{i}.docx'
      tpl.save(temp_name)
      temp_files.append(temp_name)

  return temp_files

def Surat_Daftar_Lembur(df, nama, NIK_Karyawan, bulan, url, absen_aralia, temp_file_daftar_lembur):
  # filter duplikat data
  Kegiatan = df.drop_duplicates(subset=["JangkaWaktu"]).reset_index(drop=True)
  # list tanggal lembur
  TanggalLembur = []
  for x in range(len(Kegiatan)):
    # Changed check to pd.isna
    if pd.isna(Kegiatan['TanggalAkhir'][x]):
      TanggalLembur.append(Kegiatan['TanggalAwal'][x].strftime("%d %b %Y"))
    else:
      current_date = Kegiatan['TanggalAwal'][x]
      end_date = Kegiatan['TanggalAkhir'][x]
      delta = timedelta(days=1)

      while current_date <= end_date:
          TanggalLembur.append(current_date.strftime("%d %b %Y"))
          current_date += delta

  # Table data Karyawan Form A Lembur
  Karyawan_df = pd.read_excel(url, sheet_name='Karyawan_ST') # Use a different name to avoid conflict
  Karyawan_df = Karyawan_df[Karyawan_df['NIK'] == NIK_Karyawan]
  # drop data sama
  Karyawan_df = Karyawan_df.drop_duplicates(subset=["NIK"])
  # Duplikasi data sesuai jumlah lembur
  Karyawan_df = pd.DataFrame(np.repeat(Karyawan_df.values, repeats=len(TanggalLembur), axis=0), columns=Karyawan_df.columns)
  # hapus kolom nomor surat
  Karyawan_df = Karyawan_df.iloc[:, 1:]
  # tambah satu kolom nomor baru
  Karyawan_df = Karyawan_df.reset_index().rename(columns={'index': 'No'})
  Karyawan_df['No'] = Karyawan_df['No'] + 1
  # tambah kolom tanggal lembur ke dataframe karyawan
  Karyawan_df['TanggalLembur'] = TanggalLembur
  Karyawan_df['TanggalLembur'] = pd.to_datetime(Karyawan_df['TanggalLembur'], format='%d %b %Y')

  try:
    # input jam lembur form A berdasarkan absen Aralia
    # upload file excel absen aralia dari link website ke Google Colab
    # rename menjadi report_absen.xlsx
    aralia = pd.read_excel(absen_aralia)
    # mengubah format tabel dari multi kolom menjadi satu kolom multi rows
    aralia = aralia.iloc[2:, 5:-1].transpose()
    # mengubah nama kolom
    aralia.columns = ["Jam"]
    # memisahkan baris absen datang dan absen pulang menjadi dua kolom terpisah
    Datang = aralia.iloc[0::2].reset_index().rename(columns={"Jam": "Datang", "index":"Tanggal"})
    Pulang = aralia.iloc[1::2].rename(columns={"Jam": "Pulang"}).reset_index(drop=True)
    jam_lembur = pd.merge(Datang,Pulang, left_index=True, right_index=True)

    # Ensure all entries are strings for regex operations
    jam_lembur['Datang'] = jam_lembur['Datang'].astype(str)
    jam_lembur['Pulang'] = jam_lembur['Pulang'].astype(str)

    # Extract valid time format (HH:MM:SS) from the end of the string.
    # This handles "Presensi X: HH:MM:SS" by extracting "HH:MM:SS".
    # It also handles "HH:MM:SS" directly.
    # For "Presensi X: " (without time), it will result in NaN.
    jam_lembur['Datang_cleaned'] = jam_lembur['Datang'].str.extract(r'(\d{2}:\d{2}:\d{2})$').iloc[:, 0]
    jam_lembur['Pulang_cleaned'] = jam_lembur['Pulang'].str.extract(r'(\d{2}:\d{2}:\d{2})$').iloc[:, 0]

    # Fill any NaN values (e.g., from "Presensi X: " or original NaNs) with default times.
    jam_lembur['Datang_cleaned'] = jam_lembur['Datang_cleaned'].fillna("07:30:00")
    jam_lembur['Pulang_cleaned'] = jam_lembur['Pulang_cleaned'].fillna("21:00:00")

    # Convert the cleaned time strings to datetime objects
    jam_lembur['Datang'] = pd.to_datetime(jam_lembur['Datang_cleaned'], format='%H:%M:%S')
    jam_lembur['Pulang'] = pd.to_datetime(jam_lembur['Pulang_cleaned'], format='%H:%M:%S')

    # Drop the temporary cleaned columns
    jam_lembur = jam_lembur.drop(columns=['Datang_cleaned', 'Pulang_cleaned'])

    # mengubah format tanggal dari teks ke format date
    jam_lembur['Tanggal'] = pd.to_datetime(jam_lembur['Tanggal'], format='%d/%m/%Y') # Assuming '01/04/2026' format

    # membuat kolom jam lembur
    Total_lembur = []
    for i in range(len(jam_lembur)):
      the_date = jam_lembur['Tanggal'][i]
      dtg_h = int(jam_lembur['Datang'][i].strftime('%H')) # jam kedatangan
      dtg_m = int(jam_lembur['Datang'][i].strftime('%M')) # menit kedatangan
      klr_h = int(jam_lembur['Pulang'][i].strftime('%H')) # Jam kepulangan
      klr_m = int(jam_lembur['Pulang'][i].strftime('%M')) # menit kepulangan

      # jika hari kerja maka jam lembur dimulai dari jam 5 sore
      if the_date.weekday() < 5:
        # jam kepulangan ditambah dengan hasil flooring dari menit kepulangan lalu dikurangi jam 5
        Total_lembur.append(klr_h + klr_m//30 - 17)
      else:
        # jam kepulangan ditambah dengan hasil flooring dari menit kepulangan lalu
        # dikurangi jam kedatangan ditambah dengan menit kedatangan dikurangi flooring dari menit kedatangan
        Total_lembur.append(klr_h + klr_m//30 - (dtg_h + dtg_m//30))

    # tambah kolom total jam lembur
    jam_lembur["TotalJam"] = Total_lembur

    # menggabungkan kolom jam datang, jam pulang, dan jam lembur ke dataframe utama
    Lemburan = pd.merge(Karyawan_df,jam_lembur, left_on="TanggalLembur", right_on="Tanggal", how="left")
    Lemburan = Lemburan.drop(columns=['Tanggal'])
    Lemburan['TanggalLembur'] = Lemburan['TanggalLembur'].dt.strftime('%d %b %Y')
    Lemburan['Datang'] = Lemburan['Datang'].dt.strftime('%H:%M')
    Lemburan['Pulang'] = Lemburan['Pulang'].dt.strftime('%H:%M')
  except FileNotFoundError:
    Lemburan = Karyawan_df.copy()
    Lemburan['TanggalLembur'] = Lemburan['TanggalLembur'].dt.strftime('%d %b %Y')


  #buat dictionary dari dataframe karyawan
  Lemburan_dict = Lemburan.to_dict(orient='records')

  # Template DOCX
  # doc = DocxTemplate(template_file)
  TglBuat = datetime.now().strftime("%d %B %Y")
  TglBuat = TglBuat[:3] + bulan_eng_to_ina(TglBuat[3:-5]) + TglBuat[-5:]
  lembur_dict = {
      "BulanBuat" : angka_ke_bulan(bulan),
      "TanggalBuat" : TglBuat,
      "DaftarLembur" : Lemburan_dict,
      "nama" : nama,
      "NIK_Karyawan" : NIK_Karyawan
  }

  temp_files = []
  output_file = 'daftar_lembur.docx'
  tpl = DocxTemplate(temp_file_daftar_lembur)
  tpl.render(lembur_dict)
  tpl.save(output_file)
  temp_files.append(output_file)
  return temp_files

def generate_monthly_report(Kegiatan, Karyawan, nm, bulan, absen_aralia, url,
                            temp_file_konfirmasi_absen, temp_file_perintah_lembur, temp_file_daftar_lembur):

  # nama karyawan dalam surat tugas bulan ini
  # membuat list isi karyawan
  # value satu cell yang berisi beberapa nama dipecah berdasar koma
  nama_karyawan_raw = Kegiatan['Karyawan'].fillna('').astype(str).str.split(', ')
  # kumpulan list tiap cell digbung kedalam satu list besar
  nama_karyawan = [item.strip() for sublist in nama_karyawan_raw.tolist() for item in sublist if item.strip()]
  # filter duplicate list
  nama_karyawan = list(set(nama_karyawan))

  # print("nama_karyawan")
  # print(nama_karyawan)

  # looping semua list nama karyawan lembur
  if nm in nama_karyawan:
    # filter surat tugas berdasarkan nama
    current_kegiatan = Kegiatan[Kegiatan['Karyawan'].str.contains(nm, na=False)].reset_index(drop=True)
    # jika tidak ada nama maka lanjut ke nama selanjutnya
    if current_kegiatan.empty:
      st.write(f"No activities found for {nm} in {angka_ke_bulan(bulan)}")
      return # Exit if no activities found for the employee

    # mencari NIK Karyawan berdasarkan nama
    nik_series = Karyawan.loc[Karyawan['Nama'] == nm, 'NIK']
    NIK_Karyawan = nik_series.iloc[0] if not nik_series.empty else None

    if NIK_Karyawan is None:
        st.warning(f"NIK not found for employee {nm}.")
        return

    current_month_name = datetime.now().strftime('%B')

    # membuat lokasi folder surat tugas yang telah diupload di gdrive
    # list_st = [f'/content/drive/My Drive/Python/SuratTugas/{angka_ke_bulan(bulan)}/{x}.pdf' for x in current_kegiatan['NoSurat'].unique()]

    # Konfirmasi Absen
    berkas1 = Surat_Konfirmasi_absen(current_kegiatan, nm, NIK_Karyawan, bulan, temp_file_konfirmasi_absen)
    # Surat Perintah Lembur
    berkas2 = Surat_Perintah_Lembur(current_kegiatan, nm, bulan, temp_file_perintah_lembur)
    # Surat Daftar Lembur
    berkas3 = Surat_Daftar_Lembur(current_kegiatan, nm, NIK_Karyawan, bulan, url, bulan_eng_to_ina(current_month_name), temp_file_daftar_lembur)

    # gabung tiga template
    berk = berkas1 + berkas2 + berkas3

    # nama file
    output_filename = f"{angka_ke_bulan(bulan)}_{NIK_Karyawan}_{nm}"
    absensi_output_filename = f"Absensi_{output_filename}.docx"

    # Returns the value in 'Column_B' where 'Column_A' is 'Target_Value'
    vendor_series = Karyawan.loc[Karyawan['Nama'] == nm, 'Vendor']
    vendor = vendor_series.iloc[0] if not vendor_series.empty else None

    if vendor == "EPS":
      merge_docx_files(berk[-1], berk[:-1], absensi_output_filename)
      with open(absensi_output_filename, "rb") as file:
          st.download_button(
              label="Download Surat Lembur",
              data=file,
              file_name=absensi_output_filename,
              mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          )
      st.write(absensi_output_filename)
      # time.sleep(1)
      # merge_pdfs(angka_ke_bulan(bulan), list_st, f"ST_{output_filename}.pdf")
      # print(f"ST_{output_filename}.pdf")
    else:
      # merge_pdfs(angka_ke_bulan(bulan), list_st, f"ST_{output_filename}.pdf")
      # print(f"Else_ST_{output_filename}.pdf")
      st.warning("No documents generated for merging. Check your filters.")
  else:
    st.warning(f"Employee '{nm}' not found in the activity list for the selected month.")

col1, col2 = st.columns([0.2, 0.8]) # Removed the third column
with col1:
  if st.button('Generate Surat Tugas'):
    download_ST(filtered_Kegiatan, Kegiatan, karyawan_df)
with col2:
  if st.button('Generate Surat Lembur'):
    # Make sure to pass the correct arguments to generate_monthly_report
    # 'search_name' is the employee name (nm)
    # 'selected_month_num' is the month number (bulan)
    # The remaining arguments are the template files and URL
    generate_monthly_report(filtered_Kegiatan, Karyawan, search_name, selected_month_num, "report_absen.xlsx", url,
                            temp_file_konfirmasi_absen, temp_file_perintah_lembur, temp_file_daftar_lembur)

st.markdown("**Untuk tanda tangan, pastikan hanya surat tugas yang belum selesai yang tampil di dataframe*")
st.markdown("***Jika Update Spreadsheet, lakukan 'Clear App Cache'*")
